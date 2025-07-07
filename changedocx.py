import io
import logging
from telegram import Update, File
from telegram.ext import Application, CommandHandler, MessageHandler, filters, ContextTypes
from docx import Document
from docx.shared import Pt

# --- CONFIGURATION ---
# Replace 'YOUR_TELEGRAM_BOT_TOKEN' with the token you get from BotFather
BOT_TOKEN = '8127720127:AAFeFVi4a2ZXmY-osUz9HjreJT4ZCfe4mtc'

# --- LOGGING SETUP ---
# Enables logging to see errors and bot activity in the console
logging.basicConfig(
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    level=logging.INFO
)
logger = logging.getLogger(__name__)

# --- CORE DOCUMENT MODIFICATION LOGIC ---

def set_font_size_for_cell(cell, size_in_pt):
    """Iterates through paragraphs and runs in a cell to set the font size."""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(size_in_pt)

def modify_docx(file_stream: io.BytesIO) -> io.BytesIO:
    """
    Opens a docx file from a byte stream, modifies it according to the rules,
    and returns it as a new byte stream.
    """
    try:
        # Open the document from the in-memory file stream
        document = Document(file_stream)
        logger.info(f"Document opened. Found {len(document.tables)} tables.")

        # Process each table in the document
        for table in document.tables:
            logger.info(f"Processing a table with {len(table.rows)} rows.")
            for row in table.rows:
                # Get the text from the first cell to identify the row type
                # .strip() removes any leading/trailing whitespace
                row_identifier = row.cells[0].text.strip()

                # --- APPLYING THE RULES ---

                # Rule 1: Question Text -> 14pt
                if row_identifier == 'Question':
                    target_cell = row.cells[1]
                    set_font_size_for_cell(target_cell, 14)
                    logger.info("Applied 14pt font size to 'Question' text.")

                # Rule 2: Option Text -> 12pt
                elif row_identifier == 'Option':
                    target_cell = row.cells[1]
                    set_font_size_for_cell(target_cell, 12)
                    logger.info("Applied 12pt font size to 'Option' text.")

                # Rule 3: Solution Text -> 12pt
                elif row_identifier == 'Solution':
                    target_cell = row.cells[1]
                    set_font_size_for_cell(target_cell, 12)
                    logger.info("Applied 12pt font size to 'Solution' text.")

        # Save the modified document to a new in-memory stream
        output_stream = io.BytesIO()
        document.save(output_stream)
        output_stream.seek(0) # Rewind the stream to the beginning for reading
        return output_stream

    except Exception as e:
        logger.error(f"Error processing DOCX file: {e}")
        return None


# --- TELEGRAM BOT HANDLERS ---

async def start_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handles the /start command."""
    await update.message.reply_text(
        "👋 Hello! Send me a .docx file and I will format the tables inside it for you."
    )

async def process_document_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handles receiving a .docx file."""
    if not update.message.document:
        return

    doc = update.message.document
    
    # Check if the file is a .docx file
    if doc.mime_type != 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
        await update.message.reply_text("⚠️ Please send a valid .docx file.")
        return

    await update.message.reply_text("Processing your document...")
    logger.info(f"Received file: {doc.file_name}")

    try:
        # Download the file into memory
        file_obj: File = await context.bot.get_file(doc.file_id)
        downloaded_file_stream = io.BytesIO()
        await file_obj.download_to_memory(downloaded_file_stream)
        downloaded_file_stream.seek(0)

        # Modify the document
        modified_stream = modify_docx(downloaded_file_stream)

        if modified_stream:
            # Send the modified document back
            await update.message.reply_document(
                document=modified_stream,
                filename=f"modified_{doc.file_name}"
            )
            logger.info("Successfully sent modified document back to user.")
        else:
            await update.message.reply_text("❌ Sorry, something went wrong while processing your file.")

    except Exception as e:
        logger.error(f"An error occurred in process_document_handler: {e}")
        await update.message.reply_text("❌ An unexpected error occurred.")


# --- MAIN FUNCTION TO RUN THE BOT ---

def main():
    """Starts the bot."""
    print("Bot is starting...")
    application = Application.builder().token(BOT_TOKEN).build()

    # Register handlers
    application.add_handler(CommandHandler("start", start_command))
    application.add_handler(MessageHandler(filters.Document.ALL, process_document_handler))

    # Start polling for updates
    application.run_polling()
    print("Bot has stopped.")

if __name__ == '__main__':
    main()
